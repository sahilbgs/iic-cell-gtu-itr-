import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class ReportGenerator:
    """Service to compile ActivityReport database objects into professional .docx documents."""
    
    @staticmethod
    def add_page_borders(doc):
        """Add double-line page borders on all pages of the document."""
        sectPr = doc.sections[0]._sectPr
        # XML snippet for double page borders
        pgBorders = parse_xml(
            r'<w:pgBorders %s>'
            r'  <w:top w:val="double" w:sz="12" w:space="24" w:color="auto"/>'
            r'  <w:left w:val="double" w:sz="12" w:space="24" w:color="auto"/>'
            r'  <w:bottom w:val="double" w:sz="12" w:space="24" w:color="auto"/>'
            r'  <w:right w:val="double" w:sz="12" w:space="24" w:color="auto"/>'
            r'</w:pgBorders>' % nsdecls('w')
        )
        sectPr.append(pgBorders)

    @staticmethod
    def add_header(doc, logo_path):
        """Helper to insert the centered GTU logo and university text header at the top of a page."""
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(2)
        p_logo.paragraph_format.space_before = Pt(0)
        
        # Centered circular GTU Logo
        if logo_path and os.path.exists(logo_path):
            try:
                run_logo = p_logo.add_run()
                run_logo.add_picture(logo_path, width=Inches(0.75))
            except Exception:
                # Fallback if image fails to load
                p_logo.add_run("[ GTU Logo ]").font.bold = True
        else:
            p_logo.add_run("[ GTU Logo ]").font.bold = True
            
        p_text = doc.add_paragraph()
        p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_text.paragraph_format.space_after = Pt(12)
        
        run_title = p_text.add_run("Gujarat Technological University\n")
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(12)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        
        run_naac = p_text.add_run("Accredited with A+ Grade by NAAC")
        run_naac.font.name = 'Arial'
        run_naac.font.size = Pt(10)
        run_naac.font.bold = True
        run_naac.font.color.rgb = RGBColor(0, 0, 0)

    @classmethod
    def generate(cls, report, upload_folder, logo_path):
        """
        Compiles the ActivityReport model into a styled Word document and returns it.
        
        Args:
            report (ActivityReport): The report object.
            upload_folder (str): Base upload folder path (to find photo files).
            logo_path (str): Path to gtu_logo.png file.
            
        Returns:
            docx.Document: The compiled document object.
        """
        doc = docx.Document()
        
        # 1. Page Margins Setup
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # 2. Add Double Border Page Frame
        cls.add_page_borders(doc)
        
        # 3. Add Page 1 Centered GTU Header
        cls.add_header(doc, logo_path)
        
        # 4. Report Main Headings
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(18)
        p_title.paragraph_format.space_after = Pt(6)
        
        run_t = p_title.add_run(f"Report on {report.event_type} on “{report.title}”")
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(12)
        run_t.font.bold = True
        
        p_subtitle = doc.add_paragraph()
        p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_subtitle.paragraph_format.space_after = Pt(18)
        
        run_st = p_subtitle.add_run("Gujarat Technological University – Institute of Technology and Research (GTU-ITR)")
        run_st.font.name = 'Arial'
        run_st.font.size = Pt(11)
        run_st.font.bold = True
        
        # Helper to add standard numbered list item
        def add_bullet_item(num_label, value):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(12)
            
            run_l = p.add_run(num_label)
            run_l.font.name = 'Arial'
            run_l.font.size = Pt(11)
            run_l.font.bold = True
            
            run_v = p.add_run(value)
            run_v.font.name = 'Arial'
            run_v.font.size = Pt(11)
            
        # Point 1: Title of Event
        add_bullet_item("1) Title of Event: ", report.title)
        
        # Point 2: Type of Event
        add_bullet_item("2) Type of Event: ", report.event_type)
        
        # Point 3: Date, Time and Venue (Indented lines with NO bullet points)
        p_dtv = doc.add_paragraph()
        p_dtv.paragraph_format.space_after = Pt(6)
        r_dtv = p_dtv.add_run("3) Date, Time and Venue of the Event")
        r_dtv.font.name = 'Arial'
        r_dtv.font.size = Pt(11)
        r_dtv.font.bold = True
        
        sub_fields = [
            ("Date: ", report.event_date),
            ("Time: ", report.event_time),
            ("Mode: ", report.event_mode),
            ("Venue: ", report.venue),
            ("Participants: ", report.participants_demographic)
        ]
        for lbl, val in sub_fields:
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.line_spacing = 1.15
            p_sub.paragraph_format.space_after = Pt(4)
            p_sub.paragraph_format.left_indent = Inches(0.5)
            
            rl = p_sub.add_run(lbl)
            rl.font.name = 'Arial'
            rl.font.size = Pt(11)
            rl.font.bold = True
            
            rv = p_sub.add_run(val)
            rv.font.name = 'Arial'
            rv.font.size = Pt(11)

        # Point 4: Organized By
        add_bullet_item("4) Organized By: ", report.organized_by)
        
        # Point 5: Supported By
        add_bullet_item("5) Supported By: ", report.supported_by or "N/A")
        
        # Point 6: Description of Event
        p_desc_lbl = doc.add_paragraph()
        p_desc_lbl.paragraph_format.space_after = Pt(6)
        r_desc_lbl = p_desc_lbl.add_run("6) Description of Event")
        r_desc_lbl.font.name = 'Arial'
        r_desc_lbl.font.size = Pt(11)
        r_desc_lbl.font.bold = True
        
        # Clean and write description paragraphs
        paragraphs = [p.strip() for p in report.description.split('\n') if p.strip()]
        for para_text in paragraphs:
            p_desc_p = doc.add_paragraph()
            p_desc_p.paragraph_format.line_spacing = 1.2
            p_desc_p.paragraph_format.space_after = Pt(12)
            p_desc_p.paragraph_format.left_indent = Inches(0.25)
            
            run_dp = p_desc_p.add_run(para_text)
            run_dp.font.name = 'Arial'
            run_dp.font.size = Pt(11)
            
        # Point 7: Number of Participants
        add_bullet_item("7) Number of Participants: ", str(report.num_participants))
        
        # Point 8: Geotag Photographs of the Event
        p_gp_lbl = doc.add_paragraph()
        p_gp_lbl.paragraph_format.space_after = Pt(12)
        r_gp_lbl = p_gp_lbl.add_run("8) Geotag Photographs of the Event")
        r_gp_lbl.font.name = 'Arial'
        r_gp_lbl.font.size = Pt(11)
        r_gp_lbl.font.bold = True
        
        # Load photos list from JSON
        photos = []
        if report.photos_json:
            try:
                photos = json.loads(report.photos_json)
            except Exception:
                pass
                
        # 5. Insert Photos
        # Layout strategy: stack photos vertically, approx 2 per page.
        # Since Point 8 is on the current page, if there is height we insert the first photo on this page, then break for the rest.
        # But to be safe and avoid awkward layout spacing, we can just insert the first photo under Point 8.
        # If it flows to the next page naturally, that is fine.
        # Let's stack them centered, resized to 4.5 inches width (fits 2 per page beautifully).
        
        for idx, photo in enumerate(photos):
            photo_path = os.path.join(upload_folder, photo['path'])
            if not os.path.exists(photo_path):
                continue
                
            # If it's the second image (index 1) or fourth, etc. we might want to check page breaks.
            # But the simplest approach that matches the user's PDF:
            # - Page 1 has header and items 1 to 6.
            # - Page 2 has top header, description continuation (if any), item 7, item 8, and Photo 1.
            # - Page 3 has top header, Photo 2, Photo 3.
            # To simulate this dynamically:
            # - We insert the first photo right under item 8.
            # - For all subsequent photos, if the index is odd (i.e. photo 2, 4, 6), we add a page break and insert the university header at the top of that new page.
            
            if idx > 0 and idx % 2 == 1:
                # Odd index means 2nd, 4th, 6th image -> starts a new page!
                doc.add_page_break()
                cls.add_header(doc, logo_path)
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_before = Pt(12)
                
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after = Pt(12)
            
            try:
                run_img = p_img.add_run()
                run_img.add_picture(photo_path, width=Inches(4.5))
            except Exception as e:
                p_img.add_run(f"[Error loading image: {e}]")
                
        return doc
