"""
GTU-ITR R&D & IIC Portal - Document Parser
Extracts clean text from uploaded files:
  • PDF  → pdfplumber
  • DOCX → python-docx
  • TXT / EML → plain-text read
"""

import logging
import os
import re

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Stateless helper that converts supported file formats into a single
    clean text string suitable for downstream AI processing.

    Usage::

        text = DocumentParser.parse("/path/to/file.pdf")
        text = DocumentParser.parse_bytes(raw_bytes, "application/pdf")
    """

    # Supported extensions (lowercase, without leading dot)
    SUPPORTED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt', 'eml'}

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    @classmethod
    def parse(cls, filepath: str) -> str | None:
        """
        Extract text from *filepath* based on its extension.

        Returns the cleaned text, or ``None`` on failure.
        """
        if not os.path.isfile(filepath):
            logger.error("File not found: %s", filepath)
            return None

        ext = os.path.splitext(filepath)[1].lstrip('.').lower()
        try:
            if ext == 'pdf':
                return cls._parse_pdf(filepath)
            elif ext == 'docx':
                return cls._parse_docx(filepath)
            elif ext in ('txt', 'eml'):
                return cls._parse_text(filepath)
            else:
                logger.warning("Unsupported file extension: .%s", ext)
                return None
        except Exception as exc:
            logger.exception("Failed to parse %s: %s", filepath, exc)
            return None

    @classmethod
    def parse_bytes(cls, data: bytes, filename: str) -> str | None:
        """
        Extract text from in-memory *data*.

        *filename* is used only to determine the format (via extension).
        """
        ext = os.path.splitext(filename)[1].lstrip('.').lower()
        try:
            if ext == 'pdf':
                return cls._parse_pdf_bytes(data)
            elif ext == 'docx':
                return cls._parse_docx_bytes(data)
            elif ext in ('txt', 'eml'):
                return cls._clean(data.decode('utf-8', errors='replace'))
            else:
                logger.warning("Unsupported file extension: .%s", ext)
                return None
        except Exception as exc:
            logger.exception("Failed to parse bytes (%s): %s", filename, exc)
            return None

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_pdf(filepath: str) -> str:
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return DocumentParser._clean('\n'.join(pages))

    @staticmethod
    def _parse_pdf_bytes(data: bytes) -> str:
        import io
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return DocumentParser._clean('\n'.join(pages))

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_docx(filepath: str) -> str:
        from docx import Document

        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return DocumentParser._clean('\n'.join(paragraphs))

    @staticmethod
    def _parse_docx_bytes(data: bytes) -> str:
        import io
        from docx import Document

        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return DocumentParser._clean('\n'.join(paragraphs))

    # ------------------------------------------------------------------
    # Plain text / EML
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_text(filepath: str) -> str:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as fh:
            return DocumentParser._clean(fh.read())

    # ------------------------------------------------------------------
    # Cleaning helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _clean(text: str) -> str:
        """Normalise whitespace and strip control characters."""
        # Replace form-feed, carriage-return, vertical tab with newline
        text = re.sub(r'[\x0b\x0c\r]', '\n', text)
        # Collapse 3+ newlines into 2
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Strip leading/trailing whitespace on each line
        lines = [line.strip() for line in text.splitlines()]
        text = '\n'.join(lines)
        return text.strip()
