"""
GTU-ITR R&D & IIC Portal - Principal Post Routes
Blueprint: posts  |  Prefix: /posts
"""
import os
from datetime import datetime
from flask import Blueprint, render_template, redirect, url_for, flash, request, current_app, jsonify, abort
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename

from extensions import db
from models.principal_post import PrincipalPost, POST_SOURCES, POST_STATUSES
from models.department import Department
from models.activity_report import ActivityReport
from services.report_generator import ReportGenerator
from utils.decorators import principal_required, role_required
from services.ai_post_extractor import PostExtractor


posts_bp = Blueprint('posts', __name__, url_prefix='/posts')

ALLOWED_UPLOAD_EXT = {'pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg'}


def _allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_UPLOAD_EXT


@posts_bp.route('/')
@login_required
def index():
    """List all shared principal posts for general view (excluding completed ones with reports)."""
    PrincipalPost.check_and_update_expired()
    # Filter out completed and expired posts
    posts = PrincipalPost.query.join(
        ActivityReport, PrincipalPost.id == ActivityReport.post_id, isouter=True
    ).filter(
        PrincipalPost.progress_status != 'EXPIRED',
        db.or_(
            PrincipalPost.progress_status != 'COMPLETED',
            ActivityReport.id == None,
            ActivityReport.status != 'SUBMITTED'
        )
    ).order_by(PrincipalPost.created_at.desc()).all()
    return render_template('posts/index.html', posts=posts)



@posts_bp.route('/<int:post_id>/view')
def view_post(post_id):
    """View full details and attachment of a post on a dedicated page."""
    post = PrincipalPost.query.get_or_404(post_id)
    if not current_user.is_authenticated:
        if not (post.is_public and post.approval_status == 'APPROVED'):
            return redirect(url_for('auth.login', next=request.url))
    return render_template('posts/view.html', post=post)


@posts_bp.route('/manage')
@login_required
@principal_required
def manage():
    """Management dashboard for the Principal to CRUD posts."""
    PrincipalPost.check_and_update_expired()
    posts = PrincipalPost.query.order_by(PrincipalPost.created_at.desc()).all()
    return render_template('posts/manage.html', posts=posts)


@posts_bp.route('/approved-activities')
@login_required
@principal_required
def approved_activities():
    """Principal's view of all approved activities with progress reports."""
    PrincipalPost.check_and_update_expired()
    approved_posts = PrincipalPost.query.join(
        ActivityReport, PrincipalPost.id == ActivityReport.post_id, isouter=True
    ).filter(
        PrincipalPost.approval_status == 'APPROVED',
        PrincipalPost.progress_status != 'EXPIRED',
        db.or_(
            PrincipalPost.progress_status != 'COMPLETED',
            ActivityReport.id == None,
            ActivityReport.status != 'SUBMITTED'
        )
    ).order_by(PrincipalPost.created_at.desc()).all()

    # Group by department (a post can appear in multiple departments)
    dept_groups = {}
    for post in approved_posts:
        if post.departments:
            for dept in post.departments:
                if dept.name not in dept_groups:
                    dept_groups[dept.name] = []
                dept_groups[dept.name].append(post)
        else:
            dept_name = 'Unassigned'
            if dept_name not in dept_groups:
                dept_groups[dept_name] = []
            dept_groups[dept_name].append(post)

    # Stats
    total_approved = len(approved_posts)
    completed_count = sum(1 for p in approved_posts if p.progress_status == 'COMPLETED')
    in_progress_count = sum(1 for p in approved_posts if p.progress_status == 'IN_PROGRESS')
    not_started_count = sum(1 for p in approved_posts if p.progress_status == 'NOT_STARTED')

    return render_template('posts/approved_activities.html',
                           approved_posts=approved_posts,
                           dept_groups=dept_groups,
                           total_approved=total_approved,
                           completed_count=completed_count,
                           in_progress_count=in_progress_count,
                           not_started_count=not_started_count)


def _ocr_image(image):
    """Run Tesseract OCR on a PIL Image and return extracted text."""
    import pytesseract
    return pytesseract.image_to_string(image, lang='eng')


def _extract_text_from_file(filepath):
    """Extract text from uploaded file.

    Supports:
    - **Images** (JPG/PNG/BMP/TIFF): OCR with Tesseract
    - **PDF**: Text extraction with pdfplumber; if the result is too
      short (scanned/image PDF), falls back to OCR via pypdfium2
    - **DOCX/DOC**: Extracts paragraphs + table cells + headers/footers
    - **TXT**: Plain read
    """
    import logging
    logger = logging.getLogger(__name__)

    ext = filepath.rsplit('.', 1)[1].lower()
    text = ''

    # ── Images: OCR with Tesseract ──────────────────────────────────
    if ext in ('png', 'jpg', 'jpeg', 'bmp', 'tiff', 'webp'):
        try:
            from PIL import Image
            img = Image.open(filepath)
            text = _ocr_image(img)
            logger.info("OCR extracted %d chars from image.", len(text))
        except ImportError:
            text = '[OCR unavailable – install pytesseract and Pillow]'
        except Exception as e:
            logger.warning("OCR failed for image: %s", e)
            text = f'[OCR failed: {e}]'

    # ── PDF: text first, OCR fallback for scanned pages ─────────────
    elif ext == 'pdf':
        try:
            import pdfplumber
            page_count = 0
            with pdfplumber.open(filepath) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
                    # Also extract text from tables
                    for table in (page.extract_tables() or []):
                        for row in table:
                            cells = [c.strip() for c in row if c and c.strip()]
                            if cells:
                                text += ' | '.join(cells) + '\n'

            logger.info("pdfplumber extracted %d chars from %d pages.", len(text.strip()), page_count)

            # If text extraction got very little content, the PDF is likely
            # scanned/image-based — fall back to OCR on each page
            if len(text.strip()) < 50 and page_count > 0:
                logger.info("PDF text too short (%d chars), trying OCR fallback...", len(text.strip()))
                try:
                    import pypdfium2 as pdfium
                    from PIL import Image as PILImage

                    ocr_text = ''
                    pdf_doc = pdfium.PdfDocument(filepath)
                    for i in range(min(len(pdf_doc), 5)):  # OCR max 5 pages
                        page = pdf_doc[i]
                        # Render page as image at 300 DPI for good OCR quality
                        bitmap = page.render(scale=300/72)
                        pil_image = bitmap.to_pil()
                        page_ocr = _ocr_image(pil_image)
                        if page_ocr:
                            ocr_text += page_ocr + '\n'
                    pdf_doc.close()

                    if len(ocr_text.strip()) > len(text.strip()):
                        text = ocr_text
                        logger.info("PDF OCR fallback extracted %d chars.", len(text.strip()))
                except ImportError:
                    logger.warning("pypdfium2 not available for PDF OCR fallback.")
                except Exception as ocr_e:
                    logger.warning("PDF OCR fallback failed: %s", ocr_e)

        except Exception as e:
            logger.warning("PDF extraction failed: %s", e)
            text = '[PDF text extraction failed]'

    # ── Word documents: paragraphs + tables + headers ────────────────
    elif ext in ('docx', 'doc'):
        try:
            import docx
            doc = docx.Document(filepath)
            parts = []

            # 1. Headers (often contain title / letterhead info)
            for section in doc.sections:
                header = section.header
                if header and not header.is_linked_to_previous:
                    for p in header.paragraphs:
                        if p.text.strip():
                            parts.append(p.text.strip())

            # 2. All paragraphs (main body)
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

            # 3. All tables (dates, venues, details often live here)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))

            # 4. Footers (contact info, registration links)
            for section in doc.sections:
                footer = section.footer
                if footer and not footer.is_linked_to_previous:
                    for p in footer.paragraphs:
                        if p.text.strip():
                            parts.append(p.text.strip())

            text = '\n'.join(parts)
            logger.info("DOCX extracted %d chars from %d parts.", len(text), len(parts))
        except Exception as e:
            logger.warning("DOCX extraction failed: %s", e)
            text = '[DOCX text extraction failed]'

    # ── Plain text ───────────────────────────────────────────────────
    elif ext == 'txt':
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        except Exception:
            text = '[TXT read failed]'

    return _clean_extracted_text(text)


def _clean_extracted_text(raw_text):
    """Clean up extracted text from PDF/OCR to fix common artifacts.

    Fixes:
    1. Triple-letter OCR artifacts from bold/decorative fonts
       (e.g. BBBhhhaaarrraaatttiiiyyyaaa → Bharatiya)
    2. Duplicate paragraphs / content blocks from multi-page extraction
    3. Excessive whitespace and empty lines
    """
    import re

    if not raw_text or raw_text.startswith('['):
        return raw_text.strip()

    lines = raw_text.split('\n')
    cleaned_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append('')
            continue

        # ── Fix 1: Triple-letter OCR artifact ──────────────────────
        # Detect lines like "BBBhhhaaarrraaatttiiiyyyaaa GGGyyyaaannn"
        triple_groups = re.findall(r'(.)\1{2}', stripped)

        if triple_groups:
            # Calculate what % of non-space chars are triple-repeated
            non_space = stripped.replace(' ', '')
            triple_char_count = sum(
                len(m.group(0)) for m in re.finditer(r'(.)\1{2,}', stripped)
            )
            ratio = triple_char_count / max(len(non_space), 1)

            # If more than 30% of non-space chars are in triple groups
            # AND there are at least 2 triple groups, it's an artifact
            if ratio > 0.3 and len(triple_groups) >= 2:
                fixed = re.sub(r'(.)\1{2}', r'\1', stripped)
                cleaned_lines.append(fixed)
                continue
            # Short lines that are entirely triple chars (e.g. "sssttt")
            elif ratio > 0.8 and len(stripped) <= 10:
                fixed = re.sub(r'(.)\1{2}', r'\1', stripped)
                cleaned_lines.append(fixed)
                continue

        cleaned_lines.append(stripped)

    # ── Fix 1b: Re-join broken title lines ─────────────────────────
    # Merge consecutive short lines that form a title
    # e.g. "Bharatiya Gyan Parampara: Scientific" + "Technical Development in"
    # → "Bharatiya Gyan Parampara: Scientific Technical Development in"
    merged_lines = []
    i = 0
    while i < len(cleaned_lines):
        current = cleaned_lines[i]
        # If this line ends without punctuation and next line looks like continuation
        if (current and not current.endswith(('.', ',', ':', ';', '!', '?', ''))
            and i + 1 < len(cleaned_lines)
            and cleaned_lines[i + 1]
            and not cleaned_lines[i + 1][0].isupper()
            and len(current) < 60):
            # Merge with next line
            merged_lines.append(current + ' ' + cleaned_lines[i + 1])
            i += 2
        else:
            merged_lines.append(current)
            i += 1

    # ── Fix 2: Remove duplicate lines/blocks ───────────────────────
    # Use a sliding window approach: track seen lines, skip exact repeats
    text = '\n'.join(merged_lines)

    # Split into logical blocks (separated by blank lines)
    blocks = re.split(r'\n\s*\n', text)

    seen_blocks = set()
    unique_blocks = []

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # Normalize for comparison
        key = re.sub(r'\s+', ' ', block.lower()).strip()

        # Skip if too similar to something we've seen
        if key in seen_blocks:
            continue

        # Skip if this block's content is substantially contained in an earlier block
        is_dup = False
        for seen_key in seen_blocks:
            # Check if >80% of this block's words appear in an existing block
            if len(key) > 40 and len(seen_key) > 40:
                if key in seen_key or seen_key in key:
                    is_dup = True
                    break
                # Check overlap of words
                key_words = set(key.split())
                seen_words = set(seen_key.split())
                if len(key_words) > 5:
                    overlap = len(key_words & seen_words) / len(key_words)
                    if overlap > 0.8:
                        is_dup = True
                        break
        if is_dup:
            continue

        seen_blocks.add(key)
        unique_blocks.append(block)

    # ── Fix 3: Final cleanup ───────────────────────────────────────
    result = '\n\n'.join(unique_blocks)
    result = re.sub(r'\n{3,}', '\n\n', result)

    return result.strip()


@posts_bp.route('/extract', methods=['POST'])
@login_required
@principal_required
def extract():
    """AJAX endpoint to auto-extract post details from pasted text or uploaded file.
    Supports OCR for flier images (JPG/PNG), text extraction from PDFs/DOCX,
    and saves the uploaded file as a post attachment."""
    raw_text = request.form.get('pasted_text', '').strip()
    use_ai = request.form.get('use_ai') == 'true'
    saved_attachment = None
    file_text = ''
    
    # If a file was uploaded: save as attachment + extract text from it
    if 'post_file' in request.files and request.files['post_file'].filename:
        file = request.files['post_file']
        if _allowed_file(file.filename):
            import uuid
            clean_name = secure_filename(file.filename)
            filename = f"{uuid.uuid4().hex[:8]}_{clean_name}"
            upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'posts')
            os.makedirs(upload_dir, exist_ok=True)
            filepath = os.path.join(upload_dir, filename)
            file.save(filepath)
            saved_attachment = f'posts/{filename}'
            
            # Extract text from the uploaded file (OCR for images, text for PDFs)
            file_text = _extract_text_from_file(filepath)

    # Combine: pasted text takes priority, file text is fallback
    combined_text = raw_text or file_text
            
    if not combined_text:
        if saved_attachment:
            return jsonify({
                'success': True,
                'extracted': {
                    'activity_heading': '',
                    'source': 'COMPANY',
                    'summary': '',
                    'start_date': None,
                    'end_date': None,
                    'department': None,
                    'full_content': ''
                },
                'saved_attachment': saved_attachment,
                'message': 'File attached but no text could be extracted from it.'
            })
        return jsonify({'error': 'No text or file was provided.'}), 400
        
    # Extract structured details using heuristic service
    extracted = PostExtractor.extract(combined_text, use_ai=use_ai)
    
    result = {
        'success': True,
        'extracted': extracted
    }
    if saved_attachment:
        result['saved_attachment'] = saved_attachment
    if file_text and not raw_text:
        result['ocr_used'] = True
    
    return jsonify(result)


@posts_bp.route('/create', methods=['GET', 'POST'])
@login_required
@principal_required
def create():
    """Create a new shared principal post."""
    departments = Department.query.filter_by(is_deleted=False).order_by(Department.name).all()
    
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        source = request.form.get('source', 'COMPANY')
        summary = request.form.get('summary', '').strip()
        full_content = request.form.get('full_content', '').strip()
        progress_status = request.form.get('progress_status', 'NOT_STARTED')
        department_id = request.form.get('department_id', type=int)
        
        # Validation
        if not title or not summary or not full_content:
            flash('Activity heading, summary, and full content are required.', 'danger')
            return render_template('posts/form.html', post=None, departments=departments,
                                   sources=POST_SOURCES, statuses=POST_STATUSES)
            
        # Parse Dates
        start_date = None
        start_date_str = request.form.get('start_date', '').strip()
        if start_date_str:
            try:
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            except ValueError:
                pass
                
        end_date = None
        end_date_str = request.form.get('end_date', '').strip()
        if end_date_str:
            try:
                end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
            except ValueError:
                pass

        if start_date and end_date and end_date < start_date:
            flash('End Date / Deadline cannot be earlier than Start Date.', 'danger')
            return render_template('posts/form.html', post=None, departments=departments,
                                   sources=POST_SOURCES, statuses=POST_STATUSES)

        is_public = request.form.get('is_public') == '1'
        external_registration_url = request.form.get('external_registration_url', '').strip() or None

        post = PrincipalPost(
            title=title,
            source=source,
            summary=summary,
            full_content=full_content,
            progress_status=progress_status,
            department_id=department_id or None,
            created_by=current_user.id,
            start_date=start_date,
            end_date=end_date,
            is_public=is_public,
            external_registration_url=external_registration_url
        )

        # Auto-approve if created by PRINCIPAL, CHAIRPERSON, or MASTER_ADMIN
        if current_user.role in ('PRINCIPAL', 'CHAIRPERSON', 'MASTER_ADMIN'):
            post.approval_status = 'APPROVED'
            post.approved_by = current_user.id
            post.approval_date = datetime.utcnow()

        # Link departments relationship
        if department_id:
            dept = Department.query.get(department_id)
            if dept:
                post.departments = [dept]

        # Handle file attachment
        if 'attachment' in request.files and request.files['attachment'].filename:
            file = request.files['attachment']
            if _allowed_file(file.filename):
                import uuid
                clean_name = secure_filename(file.filename)
                filename = f"{uuid.uuid4().hex[:8]}_{clean_name}"
                upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'posts')
                os.makedirs(upload_dir, exist_ok=True)
                filepath = os.path.join(upload_dir, filename)
                file.save(filepath)
                post.attachment_path = f'posts/{filename}'
            else:
                flash('Unsupported attachment file format.', 'warning')
        elif request.form.get('pre_attachment_path', '').strip():
            # Use the file that was pre-saved during AI extraction (sanitized)
            pre_path = request.form.get('pre_attachment_path').strip()
            if pre_path.startswith('posts/') and '..' not in pre_path:
                post.attachment_path = pre_path

        db.session.add(post)
        db.session.commit()
        flash('Shared Activity Post created successfully!', 'success')
        return redirect(url_for('posts.manage'))

    return render_template('posts/form.html', post=None, departments=departments,
                           sources=POST_SOURCES, statuses=POST_STATUSES)


@posts_bp.route('/<int:post_id>/edit', methods=['GET', 'POST'])
@login_required
@principal_required
def edit(post_id):
    """Edit an existing shared principal post."""
    post = PrincipalPost.query.get_or_404(post_id)
    if post.progress_status == 'COMPLETED':
        flash('Completed activities cannot be edited.', 'danger')
        return redirect(url_for('posts.manage'))
        
    departments = Department.query.filter_by(is_deleted=False).order_by(Department.name).all()
    
    if request.method == 'POST':
        post.title = request.form.get('title', '').strip()
        post.source = request.form.get('source', 'COMPANY')
        post.summary = request.form.get('summary', '').strip()
        post.full_content = request.form.get('full_content', '').strip()
        post.progress_status = request.form.get('progress_status', 'NOT_STARTED')
        post.department_id = request.form.get('department_id', type=int) or None
        
        # Validation
        if not post.title or not post.summary or not post.full_content:
            flash('Activity heading, summary, and full content are required.', 'danger')
            return render_template('posts/form.html', post=post, departments=departments,
                                   sources=POST_SOURCES, statuses=POST_STATUSES)
        
        # Parse Dates
        start_date = None
        start_date_str = request.form.get('start_date', '').strip()
        if start_date_str:
            try:
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            except ValueError:
                pass
                
        end_date = None
        end_date_str = request.form.get('end_date', '').strip()
        if end_date_str:
            try:
                end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
            except ValueError:
                pass

        if start_date and end_date and end_date < start_date:
            flash('End Date / Deadline cannot be earlier than Start Date.', 'danger')
            return render_template('posts/form.html', post=post, departments=departments,
                                   sources=POST_SOURCES, statuses=POST_STATUSES)

        post.start_date = start_date
        post.end_date = end_date
        post.is_public = request.form.get('is_public') == '1'
        post.external_registration_url = request.form.get('external_registration_url', '').strip() or None

        # Handle file attachment replacement
        if 'attachment' in request.files and request.files['attachment'].filename:
            file = request.files['attachment']
            if _allowed_file(file.filename):
                import uuid
                clean_name = secure_filename(file.filename)
                filename = f"{uuid.uuid4().hex[:8]}_{clean_name}"
                upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'posts')
                os.makedirs(upload_dir, exist_ok=True)
                filepath = os.path.join(upload_dir, filename)
                file.save(filepath)
                post.attachment_path = f'posts/{filename}'
            else:
                flash('Unsupported attachment file format.', 'warning')
        elif request.form.get('pre_attachment_path', '').strip():
            # Use the file that was pre-saved during AI extraction (sanitized)
            pre_path = request.form.get('pre_attachment_path').strip()
            if pre_path.startswith('posts/') and '..' not in pre_path:
                post.attachment_path = pre_path

        db.session.commit()
        flash('Shared Activity Post updated successfully!', 'success')
        return redirect(url_for('posts.manage'))

    return render_template('posts/form.html', post=post, departments=departments,
                           sources=POST_SOURCES, statuses=POST_STATUSES)


@posts_bp.route('/<int:post_id>/delete', methods=['POST'])
@login_required
@principal_required
def delete(post_id):
    """Delete a shared principal post."""
    post = PrincipalPost.query.get_or_404(post_id)
    if post.progress_status == 'COMPLETED':
        flash('Completed activities cannot be deleted.', 'danger')
        return redirect(url_for('posts.manage'))
        
    db.session.delete(post)
    db.session.commit()
    flash('Shared Activity Post deleted successfully.', 'success')
    return redirect(url_for('posts.manage'))


@posts_bp.route('/<int:post_id>/approve', methods=['POST'])
@login_required
@role_required('CHAIRPERSON')
def approve(post_id):
    """Chairperson approves a post and allocates it to a department."""
    post = PrincipalPost.query.get_or_404(post_id)
    department_ids = request.form.getlist('department_ids', type=int)

    if not department_ids:
        flash('Please select at least one department to allocate this activity to.', 'warning')
        return redirect(url_for('dashboard.index'))

    post.approval_status = 'APPROVED'
    post.approved_by = current_user.id
    post.approval_date = datetime.utcnow()
    post.approval_note = request.form.get('approval_note', '').strip() or None

    # Update many-to-many relationship
    depts = Department.query.filter(Department.id.in_(department_ids)).all()
    post.departments = depts

    # Maintain backward compatibility with department_id for legacy routes
    post.department_id = department_ids[0]

    db.session.commit()
    dept_names = ", ".join([d.name for d in depts])
    flash(f'Activity approved and allocated to {dept_names}!', 'success')
    return redirect(url_for('dashboard.index'))


@posts_bp.route('/<int:post_id>/reject', methods=['POST'])
@login_required
@role_required('CHAIRPERSON')
def reject(post_id):
    """Chairperson rejects a post."""
    post = PrincipalPost.query.get_or_404(post_id)
    post.approval_status = 'REJECTED'
    post.approved_by = current_user.id
    post.approval_date = datetime.utcnow()
    post.rejection_reason = request.form.get('rejection_reason', '').strip() or None

    db.session.commit()
    flash('Activity has been rejected.', 'info')
    return redirect(url_for('dashboard.index'))


@posts_bp.route('/<int:post_id>/toggle-public', methods=['POST'])
@login_required
def toggle_public(post_id):
    """Toggle public visibility for an activity on Announcements and Home Page."""
    post = PrincipalPost.query.get_or_404(post_id)
    
    if not post.can_be_managed_by(current_user):
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json:
            return jsonify({'success': False, 'error': 'You do not have permission to modify this activity’s public status.'}), 403
        flash('You do not have permission to publish or unpublish this activity.', 'danger')
        return redirect(request.referrer or url_for('posts.manage'))

    post.is_public = not post.is_public
    db.session.commit()

    status_str = "published to Public Announcements & Home Page" if post.is_public else "hidden from public view (internal only)"
    flash_msg = f'Activity "{post.title[:35]}..." {status_str}.'

    if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json:
        return jsonify({
            'success': True,
            'is_public': post.is_public,
            'message': flash_msg
        })

    flash(flash_msg, 'success')
    return redirect(request.referrer or url_for('posts.manage'))


@posts_bp.route('/uploads/<path:filename>')
def download_file(filename):
    """Serve files from the uploads directory securely."""
    from flask import send_from_directory
    as_attachment = request.args.get('download') == '1'

    # If unauthenticated, check if the file belongs to an approved, public post
    if not current_user.is_authenticated:
        base_fname = filename.split('/')[-1]
        public_post = PrincipalPost.query.filter(
            PrincipalPost.attachment_path.like(f'%{base_fname}%'),
            PrincipalPost.is_public == True,
            PrincipalPost.approval_status == 'APPROVED'
        ).first()
        if not public_post:
            return redirect(url_for('auth.login', next=request.url))

    return send_from_directory(current_app.config['UPLOAD_FOLDER'], filename, as_attachment=as_attachment)


@posts_bp.route('/<int:post_id>/assign-faculty', methods=['POST'])
@login_required
@role_required('HOD')
def assign_faculty(post_id):
    """Department Coordinator assigns a faculty lead to an approved post."""
    post = PrincipalPost.query.get_or_404(post_id)
    
    # Check if this post is allocated to the coordinator's department
    if current_user.department_id not in [d.id for d in post.departments]:
        abort(403)
        
    faculty_id = request.form.get('faculty_id', type=int)
    if not faculty_id:
        flash('Please select a valid faculty member.', 'warning')
        return redirect(url_for('dashboard.index'))
        
    # Verify the faculty belongs to coordinator's department
    from models.user import User
    faculty = User.query.filter_by(id=faculty_id, role='FACULTY', department_id=current_user.department_id, is_deleted=False).first()
    if not faculty:
        flash('Selected faculty is not in your department.', 'danger')
        return redirect(url_for('dashboard.index'))
        
    post.assigned_faculty_id = faculty_id
    db.session.commit()
    flash(f'Faculty Lead {faculty.full_name} assigned successfully!', 'success')
    return redirect(url_for('dashboard.index'))


@posts_bp.route('/<int:post_id>/update-progress', methods=['POST'])
@login_required
def update_progress(post_id):
    """Coordinator or assigned Faculty Lead updates post progress status."""
    post = PrincipalPost.query.get_or_404(post_id)
    
    # Authorization check
    is_coord = current_user.role == 'HOD' and current_user.department_id in [d.id for d in post.departments]
    is_assigned = current_user.id == post.assigned_faculty_id
    
    if not (is_coord or is_assigned):
        abort(403)
        
    progress_status = request.form.get('progress_status')
    if progress_status not in [key for key, _ in POST_STATUSES]:
        flash('Invalid progress status.', 'danger')
        return redirect(url_for('dashboard.index'))
        
    post.progress_status = progress_status
    db.session.commit()
    flash(f'Activity progress updated to {post.status_label}!', 'success')
    return redirect(url_for('dashboard.index'))


@posts_bp.route('/<int:post_id>/form-builder', methods=['GET', 'POST'])
@login_required
def form_builder(post_id):
    """Coordinator or assigned Faculty Lead accesses and configures the student registration form builder."""
    post = PrincipalPost.query.get_or_404(post_id)
    
    # Check permissions using can_be_managed_by
    if not post.can_be_managed_by(current_user):
        abort(403)
        
    import json
    from datetime import datetime
    
    # Starter suggested fields that are 100% editable, reorderable, or removable
    starter_fields = [
        {"id": "student_name", "label": "Full Name", "type": "text", "required": True, "placeholder": "Enter your full legal name"},
        {"id": "enrollment_no", "label": "Enrollment Number", "type": "text", "required": True, "placeholder": "e.g. 231040107082"},
        {"id": "email", "label": "Email Address", "type": "email", "required": True, "placeholder": "e.g. student@gtu.ac.in"},
        {"id": "phone", "label": "Phone / WhatsApp Number", "type": "tel", "required": False, "placeholder": "e.g. +91 98765 43210"},
        {"id": "department", "label": "Department", "type": "select", "options": ["Computer Engineering", "Information Technology", "Mechanical Engineering", "Civil Engineering", "Electrical Engineering", "Electronics & Communication"], "required": True},
        {"id": "semester", "label": "Current Semester", "type": "select", "options": ["1", "2", "3", "4", "5", "6", "7", "8"], "required": False}
    ]
    
    if request.method == 'POST':
        # Read form configuration from JSON input
        config_data = request.form.get('config_json', '[]')
        deadline_str = request.form.get('registration_deadline', '').strip()
        
        try:
            full_config = json.loads(config_data)
            if not isinstance(full_config, list) or len(full_config) == 0:
                full_config = starter_fields

            post.form_config = json.dumps(full_config)
            post.has_registration_form = True
            
            # Save custom registration deadline
            if deadline_str:
                try:
                    post.registration_deadline = datetime.strptime(deadline_str, '%Y-%m-%dT%H:%M')
                except ValueError:
                    try:
                        post.registration_deadline = datetime.strptime(deadline_str, '%Y-%m-%d')
                    except ValueError:
                        post.registration_deadline = None
            else:
                post.registration_deadline = None

            # Save customizable form heading, subtitle, and badge
            form_title = request.form.get('form_title', '').strip()
            form_subtitle = request.form.get('form_subtitle', '').strip()
            form_badge = request.form.get('form_badge', '').strip()

            post.form_title = form_title or None
            post.form_subtitle = form_subtitle or None
            post.form_badge = form_badge or None

            if 'external_registration_url' in request.form:
                post.external_registration_url = request.form.get('external_registration_url', '').strip() or None
            if 'is_public' in request.form:
                post.is_public = request.form.get('is_public') == '1'

            db.session.commit()
            flash('Registration form configuration and settings saved successfully!', 'success')
            return redirect(url_for('dashboard.index'))
        except Exception as e:
            flash(f'Failed to save form config: {e}', 'danger')
            
    # Load existing fields (or use starter editable fields if none set yet)
    existing_fields = []
    if post.form_config:
        try:
            existing_fields = json.loads(post.form_config)
        except Exception:
            existing_fields = []
            
    if not existing_fields:
        existing_fields = starter_fields
            
    return render_template('posts/form_builder.html', post=post, existing_fields=existing_fields)


@posts_bp.route('/<int:post_id>/form-builder/auto', methods=['POST'])
@login_required
def form_builder_auto(post_id):
    """AJAX endpoint to auto-suggest custom fields based on post details."""
    post = PrincipalPost.query.get_or_404(post_id)
    if not post.can_be_managed_by(current_user):
        return jsonify({"error": "Unauthorized"}), 403
        
    from services.form_generator import FormGenerator
    suggested_fields = FormGenerator.generate_fields(post)
    return jsonify({"success": True, "fields": suggested_fields})


@posts_bp.route('/<int:post_id>/register', methods=['GET', 'POST'])
def register(post_id):
    """Public route for students to register for the activity."""
    post = PrincipalPost.query.get_or_404(post_id)
    
    # Activity must be approved and have a registration form
    if post.approval_status != 'APPROVED' or not post.has_registration_form:
        abort(404)
        
    import json
    
    # Format closed deadline text for display
    closed_deadline_formatted = None
    if post.registration_deadline:
        closed_deadline_formatted = post.registration_deadline.strftime('%B %d, %Y at %I:%M %p')
    elif post.end_date:
        closed_deadline_formatted = post.end_date.strftime('%B %d, %Y')

    # Check if registration is closed (deadline passed or completed)
    if post.is_registration_closed:
        return render_template('posts/register.html', post=post, fields=[], 
                               is_closed=True, 
                               closed_deadline_formatted=closed_deadline_formatted or 'recently')

    # Parse form configuration
    fields = []
    if post.form_config:
        try:
            fields = json.loads(post.form_config)
        except Exception:
            pass
            
    if not fields:
        abort(500, "Registration form configuration is corrupt.")
        
    ALLOWED_REG_FILE_EXT = {'pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg', 'zip', 'rar', 'pptx', 'ppt', 'csv', 'xlsx'}

    if request.method == 'POST':
        import uuid
        import time
        from werkzeug.utils import secure_filename
        from models.student_registration import StudentRegistration

        answers = {}

        # Loop through all configured form fields
        for f in fields:
            fid = f.get('id')
            ftype = f.get('type', 'text')
            flabel = f.get('label', 'Field')
            is_req = f.get('required', False)

            val = ""
            if ftype == 'file':
                # File upload handling
                file_obj = request.files.get(f'field_{fid}') or request.files.get(fid)
                if file_obj and file_obj.filename:
                    fname = secure_filename(file_obj.filename)
                    ext = fname.rsplit('.', 1)[1].lower() if '.' in fname else ''
                    if ext not in ALLOWED_REG_FILE_EXT:
                        flash(f"File for '{flabel}' must be one of: {', '.join(sorted(ALLOWED_REG_FILE_EXT))}", 'danger')
                        return render_template('posts/register.html', post=post, fields=fields)

                    upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'registrations', str(post.id))
                    os.makedirs(upload_dir, exist_ok=True)
                    stored_fname = f"{uuid.uuid4().hex[:8]}_{fname}"
                    file_obj.save(os.path.join(upload_dir, stored_fname))
                    val = f"registrations/{post.id}/{stored_fname}"
                elif is_req:
                    flash(f"Please upload a file for '{flabel}'.", 'danger')
                    return render_template('posts/register.html', post=post, fields=fields)
            elif ftype == 'checkbox':
                vals = request.form.getlist(f'field_{fid}') or request.form.getlist(fid)
                val = ", ".join(vals)
                if is_req and not val:
                    flash(f"The field '{flabel}' is required.", 'danger')
                    return render_template('posts/register.html', post=post, fields=fields)
            else:
                val = request.form.get(f'field_{fid}', '').strip()
                if not val:
                    val = request.form.get(fid, '').strip()
                if is_req and not val:
                    flash(f"The field '{flabel}' is required.", 'danger')
                    return render_template('posts/register.html', post=post, fields=fields)

            answers[fid] = val

        # Extract core columns for StudentRegistration database record
        student_name = answers.get('student_name') or ""
        if not student_name:
            for fid, v in answers.items():
                fl = next((f.get('label', '').lower() for f in fields if f.get('id') == fid), '')
                if 'name' in fl or 'student' in fl:
                    student_name = v
                    break
        if not student_name:
            student_name = "Registered Participant"

        enrollment_no = answers.get('enrollment_no') or ""
        if not enrollment_no:
            for fid, v in answers.items():
                fl = next((f.get('label', '').lower() for f in fields if f.get('id') == fid), '')
                if 'enroll' in fl or 'roll' in fl or 'id' in fl:
                    enrollment_no = v
                    break
        if not enrollment_no:
            enrollment_no = f"REG-{int(time.time())}"

        email = answers.get('email') or ""
        if not email:
            for f in fields:
                if f.get('type') == 'email' and answers.get(f.get('id')):
                    email = answers[f['id']]
                    break

        phone = answers.get('phone') or ""
        if not phone:
            for f in fields:
                if f.get('type') == 'tel' and answers.get(f.get('id')):
                    phone = answers[f['id']]
                    break

        semester = answers.get('semester') or ""
        if not semester:
            for fid, v in answers.items():
                fl = next((f.get('label', '').lower() for f in fields if f.get('id') == fid), '')
                if 'sem' in fl:
                    semester = v
                    break

        department = answers.get('department') or ""
        if not department:
            for fid, v in answers.items():
                fl = next((f.get('label', '').lower() for f in fields if f.get('id') == fid), '')
                if 'dept' in fl or 'branch' in fl or 'department' in fl:
                    department = v
                    break

        # Duplicate check if valid enrollment or email exists
        if email or (enrollment_no and not enrollment_no.startswith('REG-')):
            dup_filters = []
            if enrollment_no and not enrollment_no.startswith('REG-'):
                dup_filters.append(StudentRegistration.enrollment_no == enrollment_no)
            if email and email != 'not-provided@gtu.ac.in':
                dup_filters.append(StudentRegistration.email == email)

            if dup_filters:
                existing_reg = StudentRegistration.query.filter(
                    StudentRegistration.post_id == post.id,
                    db.or_(*dup_filters)
                ).first()
                if existing_reg:
                    flash('You are already registered for this activity with this enrollment number or email.', 'warning')
                    return render_template('posts/register.html', post=post, fields=fields)

        # Save Student Registration
        reg = StudentRegistration(
            post_id=post.id,
            student_name=student_name,
            enrollment_no=enrollment_no,
            email=email or 'not-provided@gtu.ac.in',
            phone=phone or None,
            semester=semester or None,
            department=department or None,
            custom_data=json.dumps(answers)
        )
        db.session.add(reg)
        db.session.commit()

        # Dynamic success message
        return render_template('posts/register.html', post=post, fields=fields, success_registered=True, student_name=student_name)

    return render_template('posts/register.html', post=post, fields=fields)


@posts_bp.route('/<int:post_id>/registrations/report')
@login_required
def registration_report(post_id):
    """Detailed registrations and statistical reports, visible to Lead Faculty, Coordinator, and Higher Authorities."""
    post = PrincipalPost.query.get_or_404(post_id)

    # Authorization checks
    is_mgmt = current_user.is_management
    is_coord = current_user.role == 'HOD' and current_user.department_id in [d.id for d in post.departments]
    is_assigned = current_user.id == post.assigned_faculty_id

    if not (is_mgmt or is_coord or is_assigned):
        abort(403)

    import json

    # Parse form config to get field headers
    fields = []
    custom_field_map = {}
    if post.form_config:
        try:
            fields = json.loads(post.form_config)
            custom_field_map = {f['id']: f['label'] for f in fields}
        except Exception:
            pass

    # Fetch registrations
    from models.student_registration import StudentRegistration
    registrations = StudentRegistration.query.filter_by(post_id=post.id).order_by(StudentRegistration.registered_at.desc()).all()

    # Helper to parse custom responses in template
    def get_custom_val(reg_obj, field_id):
        if not reg_obj.custom_data:
            return ""
        try:
            answers = json.loads(reg_obj.custom_data)
            return answers.get(field_id, "")
        except Exception:
            return ""

    # Helper to get full answers dict
    def get_all_answers(reg_obj):
        if not reg_obj.custom_data:
            return {}
        try:
            return json.loads(reg_obj.custom_data)
        except Exception:
            return {}

    # Calculations / Stats
    total_regs = len(registrations)
    semester_stats = {}
    for reg in registrations:
        sem = reg.semester or "Not Specified"
        semester_stats[sem] = semester_stats.get(sem, 0) + 1

    return render_template('posts/registration_report.html',
                           post=post,
                           registrations=registrations,
                           fields=fields,
                           custom_field_map=custom_field_map,
                           total_regs=total_regs,
                           semester_stats=semester_stats,
                           get_custom_val=get_custom_val,
                           get_all_answers=get_all_answers)


@posts_bp.route('/<int:post_id>/registrations/export-csv')
@login_required
def export_registrations_csv(post_id):
    """Export all student registrations as a downloadable UTF-8 CSV/Excel spreadsheet."""
    post = PrincipalPost.query.get_or_404(post_id)

    is_mgmt = current_user.is_management
    is_coord = current_user.role == 'HOD' and current_user.department_id in [d.id for d in post.departments]
    is_assigned = current_user.id == post.assigned_faculty_id
    if not (is_mgmt or is_coord or is_assigned):
        abort(403)

    import io
    import csv
    import json
    from flask import Response
    from models.student_registration import StudentRegistration

    fields = []
    if post.form_config:
        try:
            fields = json.loads(post.form_config)
        except Exception:
            pass

    registrations = StudentRegistration.query.filter_by(post_id=post.id).order_by(StudentRegistration.registered_at.asc()).all()

    output = io.StringIO()
    output.write('\ufeff')  # UTF-8 BOM for Excel
    writer = csv.writer(output)

    # Prepare Headers
    headers = ['#', 'Student Name', 'Enrollment Number', 'Department', 'Semester', 'Email', 'Phone']
    standard_ids = {'student_name', 'enrollment_no', 'department', 'semester', 'email', 'phone'}
    for f in fields:
        if f.get('id') not in standard_ids:
            headers.append(f.get('label', f.get('id')))
    headers.append('Registered At')
    writer.writerow(headers)

    # Prepare Data Rows
    for idx, reg in enumerate(registrations, start=1):
        answers = {}
        if reg.custom_data:
            try:
                answers = json.loads(reg.custom_data)
            except Exception:
                pass

        row = [
            idx,
            reg.student_name,
            reg.enrollment_no,
            reg.department or '',
            reg.semester or '',
            reg.email or '',
            reg.phone or ''
        ]
        for f in fields:
            fid = f.get('id')
            if fid not in standard_ids:
                row.append(answers.get(fid, ''))
        row.append(reg.registered_at.strftime('%Y-%m-%d %H:%M:%S'))
        writer.writerow(row)

    safe_title = "".join(c for c in post.title[:35] if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
    filename = f"registrations_{safe_title}_{post.id}.csv"

    return Response(
        output.getvalue(),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )


@posts_bp.route('/<int:post_id>/registrations/export-excel')
@login_required
def export_registrations_excel(post_id):
    """Export all student registrations as a professionally styled Microsoft Excel (.xlsx) workbook."""
    post = PrincipalPost.query.get_or_404(post_id)

    is_mgmt = current_user.is_management
    is_coord = current_user.role == 'HOD' and current_user.department_id in [d.id for d in post.departments]
    is_assigned = current_user.id == post.assigned_faculty_id
    if not (is_mgmt or is_coord or is_assigned):
        abort(403)

    import io
    import json
    from datetime import datetime
    from flask import Response
    from models.student_registration import StudentRegistration
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    fields = []
    if post.form_config:
        try:
            fields = json.loads(post.form_config)
        except Exception:
            pass

    registrations = StudentRegistration.query.filter_by(post_id=post.id).order_by(StudentRegistration.registered_at.asc()).all()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Registrations"

    # Set grid lines visible
    ws.views.sheetView[0].showGridLines = True

    # Styling Palettes
    navy_fill = PatternFill(start_color="0F52BA", end_color="0F52BA", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    title_font = Font(name="Calibri", size=14, bold=True, color="002060")
    sub_font = Font(name="Calibri", size=10, italic=True, color="555555")
    data_font = Font(name="Calibri", size=10, color="000000")
    zebra_fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")

    thin_border_side = Side(style="thin", color="CBD5E1")
    cell_border = Border(left=thin_border_side, right=thin_border_side, top=thin_border_side, bottom=thin_border_side)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # Top Title Rows
    ws.append(["GUJARAT TECHNOLOGICAL UNIVERSITY (GTU) — ITR IIC & R&D CELL"])
    ws.cell(row=1, column=1).font = title_font
    ws.row_dimensions[1].height = 26

    depts_str = ", ".join([d.code for d in post.departments]) if post.departments else "All Departments"
    faculty_str = post.assigned_faculty.full_name if post.assigned_faculty else "Unassigned"
    ws.append([f"Activity: {post.title} | Allocated: {depts_str} | Lead Faculty: {faculty_str}"])
    ws.cell(row=2, column=1).font = sub_font
    ws.row_dimensions[2].height = 18

    ws.append([f"Total Registered Participants: {len(registrations)} | Generated: {datetime.now().strftime('%d %B %Y, %I:%M %p')}"])
    ws.cell(row=3, column=1).font = sub_font
    ws.row_dimensions[3].height = 18

    ws.append([])  # Blank row 4

    # Headers Row 5
    headers = ['#', 'Student Name', 'Enrollment Number', 'Department', 'Semester', 'Email Address', 'Phone Number']
    standard_ids = {'student_name', 'enrollment_no', 'department', 'semester', 'email', 'phone'}
    for f in fields:
        if f.get('id') not in standard_ids:
            headers.append(f.get('label', f.get('id')))
    headers.append('Registered Timestamp')

    ws.append(headers)
    header_row_idx = 5
    ws.row_dimensions[header_row_idx].height = 26

    for col_idx in range(1, len(headers) + 1):
        c = ws.cell(row=header_row_idx, column=col_idx)
        c.fill = navy_fill
        c.font = header_font
        c.alignment = center_align
        c.border = cell_border

    # Data Rows
    for r_idx, reg in enumerate(registrations, start=1):
        answers = {}
        if reg.custom_data:
            try:
                answers = json.loads(reg.custom_data)
            except Exception:
                pass

        row_data = [
            r_idx,
            reg.student_name,
            reg.enrollment_no,
            reg.department or '',
            reg.semester or '',
            reg.email or '',
            reg.phone or ''
        ]
        for f in fields:
            fid = f.get('id')
            if fid not in standard_ids:
                row_data.append(answers.get(fid, ''))
        row_data.append(reg.registered_at.strftime('%Y-%m-%d %H:%M:%S'))

        ws.append(row_data)
        curr_row = header_row_idx + r_idx
        ws.row_dimensions[curr_row].height = 22

        is_even = (r_idx % 2 == 0)
        for col_idx in range(1, len(row_data) + 1):
            cell = ws.cell(row=curr_row, column=col_idx)
            cell.font = data_font
            cell.border = cell_border
            if is_even:
                cell.fill = zebra_fill

            # Alignment
            if col_idx in (1, 3, 5, len(row_data)):
                cell.alignment = center_align
            else:
                cell.alignment = left_align

    # Auto-fit column widths
    for col in ws.columns:
        col_letter = get_column_letter(col[0].column)
        max_len = 0
        for cell in col:
            if cell.row < header_row_idx:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        ws.column_dimensions[col_letter].width = max(max_len + 4, 13) if max_len < 45 else 45

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    safe_title = "".join(c for c in post.title[:35] if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
    filename = f"registrations_{safe_title}_{post.id}.xlsx"

    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )


@posts_bp.route('/<int:post_id>/report', methods=['GET', 'POST'])
@login_required
def report_form(post_id):
    """View/edit activity report before submitting it."""
    import json
    post = PrincipalPost.query.get_or_404(post_id)
    
    # Check if user is authorized
    is_mgmt = current_user.is_management
    is_coord = current_user.role == 'HOD' and current_user.department_id in [d.id for d in post.departments]
    is_assigned = current_user.id == post.assigned_faculty_id
    
    if not (is_mgmt or is_coord or is_assigned):
        abort(403)
        
    can_edit = is_coord or is_assigned or current_user.role == 'RD_COORDINATOR'
    
    # Check if a report already exists
    report = ActivityReport.query.filter_by(post_id=post.id).first()
    
    # Prefill default values from post if no report exists
    if not report:
        from models.student_registration import StudentRegistration
        reg_count = StudentRegistration.query.filter_by(post_id=post.id).count()
        
        date_str = ""
        if post.start_date:
            date_str = post.start_date.strftime('%A, %dth %B %Y')
            if post.end_date and post.end_date != post.start_date:
                date_str += " to " + post.end_date.strftime('%A, %dth %B %Y')
        else:
            date_str = "Not Specified"
            
        report_data = {
            'title': post.title,
            'event_type': 'Online workshop' if 'online' in post.title.lower() or 'online' in post.summary.lower() else 'Workshop',
            'event_date': date_str,
            'event_time': '12:00 PM to 2:00 PM',
            'event_mode': 'Online' if 'online' in post.title.lower() or 'online' in post.summary.lower() else 'Offline',
            'venue': 'Seminar Hall, GTU-ITR, Mehsana',
            'participants_demographic': 'GTU-ITR students and faculty members',
            'organized_by': 'GTU-ITR IIC Cell & GTU Venture',
            'supported_by': 'Institution’s Innovation Council (IIC) – Ministry of Education Initiative',
            'description': post.summary + "\n\n" + post.full_content,
            'num_participants': reg_count or 0,
            'status': 'DRAFT',
            'photos': []
        }
    else:
        photos = []
        if report.photos_json:
            try:
                photos = json.loads(report.photos_json)
            except Exception:
                pass
                
        report_data = {
            'title': report.title,
            'event_type': report.event_type,
            'event_date': report.event_date,
            'event_time': report.event_time,
            'event_mode': report.event_mode,
            'venue': report.venue,
            'participants_demographic': report.participants_demographic,
            'organized_by': report.organized_by,
            'supported_by': report.supported_by,
            'description': report.description,
            'num_participants': report.num_participants,
            'status': report.status,
            'photos': photos
        }
        
    if request.method == 'POST':
        if not can_edit:
            flash('You do not have permission to modify this report.', 'danger')
            return redirect(url_for('posts.report_form', post_id=post_id))
            
        action = request.form.get('action') # 'SAVE' or 'SUBMIT'
        
        # Parse fields from form
        title = request.form.get('title', '').strip()
        event_type = request.form.get('event_type', '').strip()
        event_date = request.form.get('event_date', '').strip()
        event_time = request.form.get('event_time', '').strip()
        event_mode = request.form.get('event_mode', 'Offline').strip()
        venue = request.form.get('venue', '').strip()
        participants_demographic = request.form.get('participants_demographic', '').strip()
        organized_by = request.form.get('organized_by', '').strip()
        supported_by = request.form.get('supported_by', '').strip()
        description = request.form.get('description', '').strip()
        num_participants = request.form.get('num_participants', type=int) or 0
        
        existing_photos = []
        if report and report.photos_json:
            try:
                existing_photos = json.loads(report.photos_json)
            except Exception:
                pass
                
        updated_photos = []
        for i, photo in enumerate(existing_photos):
            if request.form.get(f'delete_photo_{i}') == 'true':
                file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], photo['path'])
                if os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except Exception:
                        pass
                continue
            caption = request.form.get(f'caption_existing_{i}', '').strip()
            updated_photos.append({
                'path': photo['path'],
                'caption': caption
            })
            
        # Handle uploads
        upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'reports', str(post.id))
        os.makedirs(upload_dir, exist_ok=True)
        
        for slot in range(6):
            file_key = f'photo_slot_{slot}'
            caption_key = f'caption_slot_{slot}'
            
            if file_key in request.files:
                file = request.files[file_key]
                caption = request.form.get(caption_key, '').strip()
                
                if file and file.filename and _allowed_file(file.filename):
                    filename = secure_filename(f"photo_{slot}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{file.filename}")
                    filepath = os.path.join(upload_dir, filename)
                    file.save(filepath)
                    updated_photos.append({
                        'path': f"reports/{post.id}/{filename}",
                        'caption': caption
                    })
                    
        # Save to DB
        if not report:
            report = ActivityReport(
                post_id=post.id,
                title=title,
                event_type=event_type,
                event_date=event_date,
                event_time=event_time,
                event_mode=event_mode,
                venue=venue,
                participants_demographic=participants_demographic,
                organized_by=organized_by,
                supported_by=supported_by,
                description=description,
                num_participants=num_participants,
                photos_json=json.dumps(updated_photos),
                status='SUBMITTED' if action == 'SUBMIT' else 'DRAFT'
            )
            db.session.add(report)
        else:
            report.title = title
            report.event_type = event_type
            report.event_date = event_date
            report.event_time = event_time
            report.event_mode = event_mode
            report.venue = venue
            report.participants_demographic = participants_demographic
            report.organized_by = organized_by
            report.supported_by = supported_by
            report.description = description
            report.num_participants = num_participants
            report.photos_json = json.dumps(updated_photos)
            
            if action == 'SUBMIT':
                report.status = 'SUBMITTED'
                
        if action == 'SUBMIT':
            post.progress_status = 'COMPLETED'
            
        db.session.commit()
        
        if action == 'SUBMIT':
            flash('Activity report has been submitted successfully, and the activity is marked COMPLETED!', 'success')
            return redirect(url_for('dashboard.index'))
        else:
            flash('Report draft saved successfully!', 'success')
            return redirect(url_for('posts.report_form', post_id=post.id))
            
    return render_template('posts/report_form.html', post=post, report=report, report_data=report_data, can_edit=can_edit)


@posts_bp.route('/<int:post_id>/report/download')
@login_required
def report_download(post_id):
    """Generate and download compiled activity report as a Word document."""
    post = PrincipalPost.query.get_or_404(post_id)
    report = ActivityReport.query.filter_by(post_id=post.id).first_or_404()
    
    is_mgmt = current_user.is_management
    is_coord = current_user.role == 'HOD' and current_user.department_id in [d.id for d in post.departments]
    is_assigned = current_user.id == post.assigned_faculty_id
    
    if not (is_mgmt or is_coord or is_assigned):
        abort(403)
        
    if report.status == 'DRAFT' and not (is_coord or is_assigned):
        abort(403, "Draft report is not visible to higher authorities yet.")
        
    logo_path = os.path.join(current_app.root_path, 'static', 'gtu_logo.png')
    
    # Generate Docx
    doc = ReportGenerator.generate(report, current_app.config['UPLOAD_FOLDER'], logo_path)
    
    export_dir = current_app.config.get('EXPORT_FOLDER', 'exports')
    os.makedirs(export_dir, exist_ok=True)
    
    filename = f"Activity_Report_{post.id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(export_dir, filename)
    doc.save(filepath)
    
    from flask import send_from_directory
    return send_from_directory(export_dir, filename, as_attachment=True)


@posts_bp.route('/completed')
@login_required
@role_required('PRINCIPAL', 'CHAIRPERSON', 'RD_COORDINATOR', 'HOD', 'FACULTY')
def completed_activities_list():
    """Page listing all completed activities with submitted reports (scoped by role)."""
    if current_user.is_management:
        completed_posts = PrincipalPost.query.join(
            ActivityReport, PrincipalPost.id == ActivityReport.post_id
        ).filter(
            PrincipalPost.progress_status == 'COMPLETED',
            ActivityReport.status == 'SUBMITTED'
        ).order_by(PrincipalPost.created_at.desc()).all()
    elif current_user.role == 'HOD' and current_user.department_id:
        completed_posts = PrincipalPost.query.join(
            ActivityReport, PrincipalPost.id == ActivityReport.post_id
        ).filter(
            PrincipalPost.progress_status == 'COMPLETED',
            ActivityReport.status == 'SUBMITTED',
            PrincipalPost.departments.any(id=current_user.department_id)
        ).order_by(PrincipalPost.created_at.desc()).all()
    elif current_user.role == 'FACULTY':
        completed_posts = PrincipalPost.query.join(
            ActivityReport, PrincipalPost.id == ActivityReport.post_id
        ).filter(
            PrincipalPost.progress_status == 'COMPLETED',
            ActivityReport.status == 'SUBMITTED',
            PrincipalPost.assigned_faculty_id == current_user.id
        ).order_by(PrincipalPost.created_at.desc()).all()
    else:
        completed_posts = []
        
    return render_template('posts/completed_activities.html', posts=completed_posts)


@posts_bp.route('/expired')
@login_required
@role_required('PRINCIPAL', 'CHAIRPERSON')
def expired_activities_list():
    """Page listing all expired activities (Principal & Chairperson only)."""
    PrincipalPost.check_and_update_expired()
    expired_posts = PrincipalPost.query.filter(
        PrincipalPost.progress_status == 'EXPIRED'
    ).order_by(PrincipalPost.created_at.desc()).all()
    
    return render_template('posts/expired_activities.html', posts=expired_posts)


@posts_bp.route('/calendar')
@login_required
def calendar_view():
    """Interactive calendar of activities based on publication/posting date."""
    PrincipalPost.check_and_update_expired()
    posts = PrincipalPost.query.filter(
        PrincipalPost.progress_status != 'EXPIRED'
    ).all()
    return render_template('posts/calendar.html', posts=posts)

